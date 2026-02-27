# Copyright (c) 2025-2026 Filippo Danesi. All rights reserved.
"""DOCX writer — inserts hyperlinks via python-docx XML manipulation."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from seo_linker.models import LinkingResult
from seo_linker.writers.base import BaseWriter

# Pattern to find markdown links
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
# Pattern to detect markdown headings
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)")
# Pattern to detect **bold** spans
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
# Markdown table separator row (e.g. |---|---|)
TABLE_SEP_RE = re.compile(r"^\|[\s\-:|]+\|$")


class DocxWriter(BaseWriter):
    def write(self, result: LinkingResult, input_path: Path, output_path: Path) -> None:
        if result.rewritten_text:
            # Rewrite path: paragraphs changed, so rebuild the doc from linked_text
            doc = _write_from_linked_text(result, input_path)
        else:
            # No-rewrite path: match original paragraphs in the .docx
            doc = Document(str(input_path))
            link_map = _build_link_map(result)
            for para in doc.paragraphs:
                plain = para.text.strip()
                if plain in link_map:
                    _replace_paragraph_with_links(para, link_map[plain])

        # Insert SEO metadata block at the end of the document
        if result.seo_title or result.seo_meta_description:
            _insert_seo_metadata(doc, result.seo_title, result.seo_meta_description)

        doc.save(str(output_path))


def _write_from_linked_text(result: LinkingResult, input_path: Path) -> Document:
    """Build a new Document from linked_text when rewrite was active.

    Copies base font info from the original .docx, then writes headings,
    body paragraphs (with hyperlinks), and plain-text table rows.
    """
    orig_doc = Document(str(input_path))

    # Extract base font info from the first body paragraph in the original doc
    base_font_name = None
    base_font_size = None
    for para in orig_doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            continue
        for run in para.runs:
            if run.font.name:
                base_font_name = run.font.name
            if run.font.size:
                base_font_size = run.font.size
            if base_font_name and base_font_size:
                break
        if base_font_name and base_font_size:
            break

    doc = Document()
    lines = result.linked_text.split("\n")
    i = 0

    while i < len(lines):
        stripped = lines[i].strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Heading lines
        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            level = len(heading_match.group(1))  # 1-6
            text = heading_match.group(2).strip()
            doc.add_heading(text, level=min(level, 9))
            i += 1
            continue

        # Markdown table — collect consecutive | lines and build a real table
        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            _add_markdown_table(doc, table_lines, base_font_name)
            continue

        # Bullet list items (- text or * text)
        if re.match(r"^[-*]\s+", stripped):
            text = re.sub(r"^[-*]\s+", "", stripped)
            para = doc.add_paragraph(style="List Bullet")
            _populate_paragraph_with_links(
                para, text, base_font_name, base_font_size
            )
            i += 1
            continue

        # Normal paragraph — parse markdown links + bold into runs
        para = doc.add_paragraph()
        _populate_paragraph_with_links(para, stripped, base_font_name, base_font_size)
        i += 1

    return doc


def _populate_paragraph_with_links(
    para, text: str, font_name: str | None, font_size: Pt | None
) -> None:
    """Parse markdown links and **bold** in *text* and populate *para*."""
    last_end = 0
    for match in LINK_RE.finditer(text):
        before = text[last_end : match.start()]
        if before:
            _add_formatted_runs(para, before, font_name, font_size)

        _add_hyperlink(para, match.group(1), match.group(2))
        last_end = match.end()

    remaining = text[last_end:]
    if remaining:
        _add_formatted_runs(para, remaining, font_name, font_size)


def _add_formatted_runs(
    para, text: str, font_name: str | None, font_size: Pt | None
) -> None:
    """Add runs to *para*, converting **bold** markers to bold formatting."""
    last = 0
    for match in BOLD_RE.finditer(text):
        # Text before the bold span
        before = text[last : match.start()]
        if before:
            run = para.add_run(before)
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = font_size
        # Bold text
        run = para.add_run(match.group(1))
        run.bold = True
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        last = match.end()

    # Text after last bold span (or entire text if no bold)
    after = text[last:]
    if after:
        run = para.add_run(after)
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size


def _add_markdown_table(doc: Document, table_lines: list[str], font_name: str | None) -> None:
    """Parse markdown table lines and add a real docx table."""
    # Split cells from each row, stripping outer pipes
    def _parse_row(line: str) -> list[str]:
        cells = line.strip().strip("|").split("|")
        return [c.strip() for c in cells]

    # Filter out separator rows (|---|---|)
    data_rows = [ln for ln in table_lines if not TABLE_SEP_RE.match(ln)]
    if not data_rows:
        return

    parsed = [_parse_row(r) for r in data_rows]
    n_cols = max(len(r) for r in parsed)
    n_rows = len(parsed)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = "Table Grid"

    for row_idx, cells in enumerate(parsed):
        for col_idx, cell_text in enumerate(cells):
            if col_idx >= n_cols:
                break
            cell = table.rows[row_idx].cells[col_idx]
            para = cell.paragraphs[0]
            # Parse bold in cell text
            _add_formatted_runs(para, cell_text, font_name, None)
            # Bold the header row
            if row_idx == 0:
                for run in para.runs:
                    run.bold = True


def _build_link_map(result: LinkingResult) -> dict[str, str]:
    """Map original paragraph text to linked version for paragraphs that changed."""
    orig_paras = result.original_text.split("\n")
    linked_paras = result.linked_text.split("\n")

    mapping: dict[str, str] = {}

    # Try to match paragraphs by position
    orig_clean = [p.strip() for p in orig_paras if p.strip()]
    linked_clean = [p.strip() for p in linked_paras if p.strip()]

    for orig, linked in zip(orig_clean, linked_clean):
        if orig != linked and LINK_RE.search(linked):
            # Strip any existing markdown links from orig for matching
            orig_no_links = LINK_RE.sub(r"\1", orig)
            mapping[orig_no_links] = linked
            mapping[orig] = linked

    return mapping


def _replace_paragraph_with_links(para, linked_text: str) -> None:
    """Replace paragraph content with text containing hyperlinks."""
    # Clear existing runs
    for run in para.runs:
        run.text = ""

    # Remove all existing run and hyperlink elements
    for child in list(para._element):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            para._element.remove(child)

    # Parse the linked text and insert runs + hyperlinks
    last_end = 0
    for match in LINK_RE.finditer(linked_text):
        # Add text before link
        before = linked_text[last_end : match.start()]
        if before:
            _add_run(para, before)

        anchor = match.group(1)
        url = match.group(2)
        _add_hyperlink(para, anchor, url)
        last_end = match.end()

    # Add remaining text
    remaining = linked_text[last_end:]
    if remaining:
        _add_run(para, remaining)


def _add_run(para, text: str) -> None:
    run_el = OxmlElement("w:r")
    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    run_el.append(text_el)
    para._element.append(run_el)


def _add_hyperlink(para, anchor_text: str, url: str) -> None:
    """Add a hyperlink to a paragraph with explicit blue + underline styling."""
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Style reference (works when the style exists in the doc)
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)

    # Explicit blue color — ensures visibility even without the Hyperlink style
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    # Explicit underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    run.append(rPr)

    text_el = OxmlElement("w:t")
    text_el.text = anchor_text
    text_el.set(qn("xml:space"), "preserve")
    run.append(text_el)

    hyperlink.append(run)
    para._element.append(hyperlink)


def _insert_seo_metadata(doc: Document, title: str, meta_description: str) -> None:
    """Insert a styled SEO metadata block at the end of the document."""
    # Separator line before the block
    sep = doc.add_paragraph()
    sep_pPr = OxmlElement("w:pPr")
    sep_pBdr = OxmlElement("w:pBdr")
    top_border = OxmlElement("w:top")
    top_border.set(qn("w:val"), "single")
    top_border.set(qn("w:sz"), "4")
    top_border.set(qn("w:space"), "1")
    top_border.set(qn("w:color"), "999999")
    sep_pBdr.append(top_border)
    sep_pPr.append(sep_pBdr)
    sep._element.insert(0, sep_pPr)

    # Header line
    header_p = doc.add_paragraph()
    header_run = header_p.add_run("SEO METADATA")
    header_run.bold = True
    header_run.font.size = Pt(10)
    header_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"Title: {title}")
    title_run.font.size = Pt(9)
    title_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Meta description
    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(f"Meta Description: {meta_description}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
